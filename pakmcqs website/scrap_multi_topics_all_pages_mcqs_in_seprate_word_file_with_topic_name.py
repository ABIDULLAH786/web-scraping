import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from civil_eng_multi_topics import multi_topics  # Importing the topics list

def sanitize_filename(filename):
    """Replace invalid characters in filenames with underscores."""
    return re.sub(r'[\/:*?"<>|]', '_', filename)

def scrape_mcqs(url, start_idx, topic_title):
    page = requests.get(url)
    soup = BeautifulSoup(page.text, 'html.parser')
    if soup.title.text == "Page not found - PakMcqs":
        print("Page title: ", soup.title.text)
        return None, start_idx

    all_mcqs = soup.find_all("article", class_="l-post grid-post grid-card-post")
    mcq_list = []
    for mcq in all_mcqs:
        item = {}
        item["#Sr."] = start_idx
        item["topic"] = topic_title
        item["question"] = mcq.find('h2').text.strip()
        
        options_container = mcq.find("div", class_="excerpt")
        options_html = options_container.find('p')
        options = options_html.get_text(separator="\n").strip().split('\n')
        options = [opt.strip() for opt in options if opt.strip()]  # Remove empty options and strip extra spaces
        options = [opt[3:].strip() if len(opt) > 1 and opt[1] == '.' else opt for opt in options]  # Remove option labels

        # Ignore the options if the correct one is not found and place the cell empty
        correct_option_text = ""
        if options_html.find('strong') is not None:
            correct_option_text = options_html.find('strong').text.strip()
            correct_option_text = correct_option_text[3:].strip() if len(correct_option_text) > 1 and correct_option_text[1] == '.' else correct_option_text  # Remove option labels from correct option

        # Assign options to respective columns
        item["A"] = options[0] if len(options) > 0 else ""
        item["B"] = options[1] if len(options) > 1 else ""
        item["C"] = options[2] if len(options) > 2 else ""
        item["D"] = options[3] if len(options) > 3 else ""

        # Determine the correct option letter
        if correct_option_text == item["A"]:
            item["Ans"] = "A"
        elif correct_option_text == item["B"]:
            item["Ans"] = "B"
        elif correct_option_text == item["C"]:
            item["Ans"] = "C"
        elif correct_option_text == item["D"]:
            item["Ans"] = "D"
        else:
            item["Ans"] = ""

        if item["Ans"]:  # Only save MCQs that have an answer
            mcq_list.append(item)
        start_idx += 1
    
    return mcq_list, start_idx

def set_font(para, font_name='Arial', font_size=12, bold=False):
    """Set font for the paragraph."""
    run = para.add_run()
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    return run

start_idx = 1

for topic in multi_topics:
    print(f"=> {topic['id']} Extracting the topic: {topic['title']}")
    current_page = 1
    proceed = True
    topic_data = []
    
    while proceed:
        print(f"\t\t - Scraping page {current_page} of topic '{topic['title']}'")
        url = topic['url'].replace('page/1', f'page/{current_page}')
        print(f"\t\t\t - Extracting URL: {url}")
        
        mcqs_data, start_idx = scrape_mcqs(url, start_idx, topic['title'])
        if mcqs_data is None:
            proceed = False
        else:
            print(f"\t\t\t - Total extracted MCQs from page {current_page}: {len(mcqs_data)}")
            topic_data.extend(mcqs_data)
            current_page += 1
        
        print(f"\t\t\t - Overall MCQs for topic '{topic['title']}': {len(topic_data)}")

    if topic_data:
        # Create a new Word document for each topic
        doc = Document()
        
        # Set two-column layout
        section = doc.sections[0]
        section._sectPr = section._sectPr.xpath('//w:cols')[0]
        section._sectPr[0].set("num", "2")
        
        # Add the topic title in bold
        heading = doc.add_paragraph()
        set_font(heading, font_size=16, bold=True).add_run(topic['title'])
        doc.add_paragraph("\n")

        # Loop through the MCQs and add them to the document
        for item in topic_data:
            # Add question in bold and font 13px
            question = doc.add_paragraph()
            set_font(question, font_size=13, bold=True).add_run(item['question'])
            
            # Add the options
            options_para = doc.add_paragraph()
            for option, label in zip([item["A"], item["B"], item["C"], item["D"]], ["A", "B", "C", "D"]):
                option_text = f"{label}. {option}"
                if option == item["Ans"]:  # Highlight the correct option
                    set_font(options_para, font_size=12, bold=True).add_run(option_text)
                else:
                    set_font(options_para, font_size=12).add_run(option_text)
                options_para.add_run("\n")  # Add a new line between options

            doc.add_paragraph("\n")  # Add space between MCQs

        # Save the file with a sanitized topic name
        safe_topic_name = sanitize_filename(topic['title'])
        file_name = f"{safe_topic_name}.docx"
        doc.save(file_name)
        print(f"Saved {file_name}")

print("Scraping complete. All topics saved separately.")
