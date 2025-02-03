from bs4 import BeautifulSoup
import requests
from docx import Document
import re
import os

# Base URL for the course
base_url = "https://www.programiz.com"
course_url = "/python-programming"

# Make a request to fetch the HTML content of the main course page
response = requests.get(base_url + course_url)
if response.status_code != 200:
    print(f"Failed to fetch the page: {response.status_code}")
    exit()

# Parse the HTML using BeautifulSoup
soup = BeautifulSoup(response.content, 'html.parser')

# Container for the course data
course_data = []

# Find the main container with the specified classes
main_container = soup.find('div', class_='contents landing_ad-wrapper')
if not main_container:
    print("The specified container was not found on the page.")
    exit()

# Find all accordion nodes within the container
accordion_items = main_container.find_all('div', class_='accordion__nodes')

# Iterate over each accordion item to extract data
for item in accordion_items:
    # Extract the section title
    header = item.find('div', class_='accordion-header__title')
    section_title = header.get_text(strip=True) if header else ""

    # Extract all links under this section
    links = item.find_all('a')
    lessons = [
        {
            'title': link.get_text(strip=True),
            'url': link['href']
        }
        for link in links
    ]

    # Add to course data
    course_data.append({
        'section': section_title,
        'lessons': lessons
    })

# Function to clean invalid characters from filenames
def clean_filename(filename):
    # Remove characters that are not valid in filenames
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

# Export each topic to a Word file
for section in course_data:
    doc = Document()
    doc.add_heading(section['section'], level=1)

    for lesson in section['lessons']:
        lesson_url = base_url + lesson['url'] if not lesson['url'].startswith('http') else lesson['url']
        lesson_response = requests.get(lesson_url)
        
        if lesson_response.status_code != 200:
            print(f"Failed to fetch lesson: {lesson['title']} ({lesson_url})")
            continue

        # Parse the HTML of the lesson page
        lesson_soup = BeautifulSoup(lesson_response.content, 'html.parser')

        # Find the content of the lesson inside the editor-contents class
        editor_contents = lesson_soup.find('div', class_='editor-contents')
        if editor_contents:
            heading = editor_contents.find('h1').get_text(strip=True)
            content = editor_contents.get_text("\n", strip=True)
            
            # Add heading and content to the document
            doc.add_heading(heading, level=2)
            doc.add_paragraph(content)

        # Clean the filename and save the Word document
        filename = f"{section['section'].replace(' ', '_')}_{lesson['title'].replace(' ', '_')}.docx"
        filename = clean_filename(filename)
        
        try:
            doc.save(filename)
            print(f"Document saved: {filename}")
        except Exception as e:
            print(f"Failed to save document {filename}: {e}")

    # Print a message indicating success for the section
    print(f"Word document created for section: {section['section']}")
